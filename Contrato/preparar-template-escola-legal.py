"""Gera contrato-escola-legal-template.docx a partir do contrato oficial (tags docxtemplater)."""
from docx import Document
from docx.shared import Inches
from pathlib import Path

ORIGEM = Path(__file__).parent / "Contrato - Projeto ESCOLA LEGAL.docx"
SAIDA = Path(__file__).parent / "contrato-escola-legal-template.docx"
ASSINATURA_DOUTORA = Path(__file__).parent / "Assinatura Doutora.jpeg"

CAMPOS = [
    (5, "Instituição: {instituicao}"),
    (6, "CNPJ nº. {cnpj}"),
    (7, "Endereço: {endereco_linha1}"),
    (8, "{endereco_linha2}"),
    (9, "Representante: {representante}"),
    (10, "CPF nº. {cpf}"),
]


def definir_texto(paragraph, texto):
    if paragraph.runs:
        paragraph.runs[0].text = texto
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(texto)


def main():
    doc = Document(ORIGEM)
    for idx, texto in CAMPOS:
        if idx < len(doc.paragraphs):
            definir_texto(doc.paragraphs[idx], texto)

    if len(doc.paragraphs) > 11:
        definir_texto(doc.paragraphs[11], "{%assinatura_contratante}")

    final = doc.add_paragraph()
    definir_texto(final, "{%assinatura_contratante_final}")

    if ASSINATURA_DOUTORA.exists():
        p_doutora = doc.add_paragraph()
        p_doutora.add_run().add_picture(str(ASSINATURA_DOUTORA), width=Inches(2.2))

    doc.save(SAIDA)
    print(f"OK - template em {SAIDA}")


if __name__ == "__main__":
    main()
